from io import BytesIO
from typing import Optional, Union, overload

from docx import Document
from faker.providers import BaseProvider

from ..base import BytesValue, DynamicTemplate, FileMixin, StringValue
from ..constants import DEFAULT_TEXT_MAX_NB_CHARS
from ..storages.base import BaseStorage
from ..storages.filesystem import FileSystemStorage

__author__ = "Artur Barseghyan <artur.barseghyan@gmail.com>"
__copyright__ = "2022-2023 Artur Barseghyan"
__license__ = "MIT"
__all__ = ("DocxFileProvider",)


class DocxFileProvider(BaseProvider, FileMixin):
    """DOCX file provider.

    Usage example:

        from faker import Faker
        from faker_file.providers.docx_file import DocxFileProvider

        file = DocxFileProvider(Faker()).docx_file()

    Usage example with options:

        file = DocxFileProvider(Faker()).docx_file(
            prefix="zzz",
            max_nb_chars=100_000,
            wrap_chars_after=80,
        )

    Usage example with `FileSystemStorage` storage (for `Django`):

        from django.conf import settings
        from faker_file.storages.filesystem import FileSystemStorage

        file = DocxFileProvider(Faker()).docx_file(
            storage=FileSystemStorage(
                root_path=settings.MEDIA_ROOT,
                rel_path="tmp",
            ),
            prefix="zzz",
            max_nb_chars=100_000,
            wrap_chars_after=80,
        )

    Usage example with content modifiers:

        from io import BytesIO
        from faker_file.base import DynamicTemplate
        from faker_file.providers.jpeg_file import JpegFileProvider

        def add_table(provider, document, data, counter, **kwargs):
            table = document.add_table(
                kwargs.get("rows", 3),
                kwargs.get("cols", 4),
            )
            data.setdefault("content_modifiers", {})
            data["content_modifiers"].setdefault("add_table", {})
            data["content_modifiers"]["add_table"].setdefault(counter, [])

            for row in table.rows:
                for cell in row.cells:
                    text = provider.generator.paragraph()
                    cell.text = text
                    data["content_modifiers"]["add_table"][counter].append(
                        text
                    )
                    data["content"] += ("\r\n" + text)


        def add_picture(provider, document, data, counter, **kwargs):
            jpeg_file = JpegFileProvider(provider.generator).jpeg_file(
                raw=True
            )
            picture = document.add_picture(BytesIO(jpeg_file))

            data.setdefault("content_modifiers", {})
            data["content_modifiers"].setdefault("add_picture", {})
            data["content_modifiers"]["add_picture"].setdefault(counter, [])

            data["content_modifiers"]["add_picture"][counter].append(
                jpeg_file.data["content"]
            )
            data["content"] += ("\r\n" + jpeg_file.data["content"])


        file = DocxFileProvider(Faker()).docx_file(
            content=DynamicTemplate([(add_table, {}), (add_picture, {})])
        )
    """

    extension: str = "docx"

    @overload
    def docx_file(
        self: "DocxFileProvider",
        storage: Optional[BaseStorage] = None,
        basename: Optional[str] = None,
        prefix: Optional[str] = None,
        max_nb_chars: int = DEFAULT_TEXT_MAX_NB_CHARS,
        wrap_chars_after: Optional[int] = None,
        content: Optional[Union[str, DynamicTemplate]] = None,
        raw: bool = True,
        **kwargs,
    ) -> BytesValue:
        ...

    @overload
    def docx_file(
        self: "DocxFileProvider",
        storage: Optional[BaseStorage] = None,
        basename: Optional[str] = None,
        prefix: Optional[str] = None,
        max_nb_chars: int = DEFAULT_TEXT_MAX_NB_CHARS,
        wrap_chars_after: Optional[int] = None,
        content: Optional[Union[str, DynamicTemplate]] = None,
        **kwargs,
    ) -> StringValue:
        ...

    def docx_file(
        self: "DocxFileProvider",
        storage: Optional[BaseStorage] = None,
        basename: Optional[str] = None,
        prefix: Optional[str] = None,
        max_nb_chars: int = DEFAULT_TEXT_MAX_NB_CHARS,
        wrap_chars_after: Optional[int] = None,
        content: Optional[Union[str, DynamicTemplate]] = None,
        raw: bool = False,
        **kwargs,
    ) -> Union[BytesValue, StringValue]:
        """Generate a DOCX file with random text.

        :param storage: Storage. Defaults to `FileSystemStorage`.
        :param basename: File basename (without extension).
        :param prefix: File name prefix.
        :param max_nb_chars: Max number of chars for the content.
        :param wrap_chars_after: If given, the output string would be separated
             by line breaks after the given position.
        :param content: File content. Might contain dynamic elements (still
            being a string), which are then replaced by correspondent
            fixtures. Can alternatively be a `DynamicTemplate` - list of content
            modifiers (callables to call after the document instance has been
            created). Each callable should accept the following
            arguments: provider, document, data, counter and **kwargs.
        :param raw: If set to True, return `BytesValue` (binary content of
            the file). Otherwise, return `StringValue` (path to the saved
            file).
        :return: Relative path (from root directory) of the generated file
            or raw content of the file.
        """
        # Generic
        if storage is None:
            storage = FileSystemStorage()

        filename = storage.generate_filename(
            extension=self.extension,
            prefix=prefix,
            basename=basename,
        )

        if isinstance(content, DynamicTemplate):
            _content = ""
        else:
            _content = self._generate_text_content(
                max_nb_chars=max_nb_chars,
                wrap_chars_after=wrap_chars_after,
                content=content,
            )
        data = {"content": _content, "filename": filename}

        with BytesIO() as _fake_file:
            document = Document()
            if _content:
                document.add_paragraph(_content)
            elif isinstance(content, DynamicTemplate):
                for counter, (ct_modifier, ct_modifier_kwargs) in enumerate(
                    content.content_modifiers
                ):
                    ct_modifier(
                        self,
                        document,
                        data,
                        counter,
                        **ct_modifier_kwargs,
                    )

            document.save(_fake_file)

            if raw:
                raw_content = BytesValue(_fake_file.getvalue())
                raw_content.data = data
                return raw_content

            storage.write_bytes(filename, _fake_file.getvalue())

        # Generic
        file_name = StringValue(storage.relpath(filename))
        file_name.data = data
        return file_name
