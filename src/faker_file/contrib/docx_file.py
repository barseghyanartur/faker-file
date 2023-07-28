from io import BytesIO

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from ..base import DEFAULT_FORMAT_FUNC

__author__ = "Artur Barseghyan <artur.barseghyan@gmail.com>"
__copyright__ = "2022-2023 Artur Barseghyan"
__license__ = "MIT"
__all__ = (
    "add_h1_heading",
    "add_h2_heading",
    "add_h3_heading",
    "add_h4_heading",
    "add_h5_heading",
    "add_h6_heading",
    "add_heading",
    "add_page_break",
    "add_paragraph",
    "add_picture",
    "add_table",
    "add_title_heading",
)


def add_table(provider, document, data, counter, **kwargs):
    """Callable responsible for the table generation."""
    table = document.add_table(
        kwargs.get("rows", 3),
        kwargs.get("cols", 4),
    )

    # Modifications of `data` is not required for generation
    # of the file, but is useful for when you want to get
    # the text content of the file.
    data.setdefault("content_modifiers", {})
    data["content_modifiers"].setdefault("add_table", {})
    data["content_modifiers"]["add_table"].setdefault(counter, [])

    for row in table.rows:
        for cell in row.cells:
            text = provider.generator.paragraph()
            cell.text = text
            # Useful when you want to get the text content of the file.
            data["content_modifiers"]["add_table"][counter].append(text)
            data["content"] += "\r\n" + text


def add_picture(provider, document, data, counter, **kwargs):
    """Callable responsible for the picture generation."""
    image = kwargs.get("image", provider.generator.image())
    document.add_picture(BytesIO(image))

    # # Modifications of `data` is not required for generation
    # # of the file, but is useful for when you want to get
    # # the text content of the file.
    # data.setdefault("content_modifiers", {})
    # data["content_modifiers"].setdefault("add_picture", {})
    # data["content_modifiers"]["add_picture"].setdefault(counter, [])
    # data["content_modifiers"]["add_picture"][counter].append(
    #     jpeg_file.data["content"]
    # )
    # data["content"] += "\r\n" + jpeg_file.data["content"]


def add_page_break(provider, document, data, counter, **kwargs):
    """Callable responsible for page break generation."""
    # Insert a page break
    document.add_page_break()


def add_paragraph(provider, document, data, counter, **kwargs):
    """Callable responsible for the paragraph generation."""
    content = kwargs.get("content", None)
    max_nb_chars = kwargs.get("content", 5_000)
    wrap_chars_after = kwargs.get("wrap_chars_after", None)
    format_func = kwargs.get("format_func", DEFAULT_FORMAT_FUNC)

    _content = provider._generate_text_content(
        max_nb_chars=max_nb_chars,
        wrap_chars_after=wrap_chars_after,
        content=content,
        format_func=format_func,
    )
    document.add_paragraph(_content)

    # Meta-data
    data.setdefault("content_modifiers", {})
    data["content_modifiers"].setdefault("add_paragraph", {})
    data["content_modifiers"]["add_paragraph"].setdefault(counter, [])
    data["content_modifiers"]["add_paragraph"][counter].append(_content)
    data["content"] += "\r\n" + _content


def add_heading(provider, document, data, counter, **kwargs):
    """Callable responsible for the heading generation."""
    content = kwargs.get("content", None)
    max_nb_chars = kwargs.get("content", 30)
    wrap_chars_after = kwargs.get("wrap_chars_after", None)
    format_func = kwargs.get("format_func", DEFAULT_FORMAT_FUNC)
    level = kwargs.get("level", 0)

    _content = provider._generate_text_content(
        max_nb_chars=max_nb_chars,
        wrap_chars_after=wrap_chars_after,
        content=content,
        format_func=format_func,
    )

    heading = document.add_heading(level=level)
    heading.add_run(_content)

    # Aligning the heading to center
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Meta-data
    data.setdefault("content_modifiers", {})
    data["content_modifiers"].setdefault("add_heading", {})
    data["content_modifiers"]["add_heading"].setdefault(counter, [])
    data["content_modifiers"]["add_heading"][counter].append(_content)
    data["content"] += "\r\n" + _content


def add_title_heading(provider, document, data, counter, **kwargs):
    """Callable responsible for the title heading generation."""
    return add_heading(provider, document, data, counter, level=0, **kwargs)


def add_h1_heading(provider, document, data, counter, **kwargs):
    """Callable responsible for the h1 heading generation."""
    return add_heading(provider, document, data, counter, level=1, **kwargs)


def add_h2_heading(provider, document, data, counter, **kwargs):
    """Callable responsible for the h2 heading generation."""
    return add_heading(provider, document, data, counter, level=2, **kwargs)


def add_h3_heading(provider, document, data, counter, **kwargs):
    """Callable responsible for the h3 heading generation."""
    return add_heading(provider, document, data, counter, level=3, **kwargs)


def add_h4_heading(provider, document, data, counter, **kwargs):
    """Callable responsible for the h4 heading generation."""
    return add_heading(provider, document, data, counter, level=4, **kwargs)


def add_h5_heading(provider, document, data, counter, **kwargs):
    """Callable responsible for the h5 heading generation."""
    return add_heading(provider, document, data, counter, level=5, **kwargs)


def add_h6_heading(provider, document, data, counter, **kwargs):
    """Callable responsible for the h6 heading generation."""
    return add_heading(provider, document, data, counter, level=6, **kwargs)
